import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Load the data from the csv file
data = pd.read_csv('assets/certdata.csv')

# Initialize Table
table = doc.add_table(rows=7, cols=2)

# Add the column names to the table
for i, column in enumerate(data.columns):
    table.cell(0, i).text = column

# Add the data to the table
for index, row in data.iterrows():
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = str(value)

# define images
company_image = doc.add_picture('assets/cert-border.png')
# Add the second image
background_image = doc.add_picture('assets/cert-bmk.png')

# Set the document to landscape orientation
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

company_image = doc.paragraphs[0]
company_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

background_image = doc.paragraphs[1]
background_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a paragraph between the images
#doc.add_paragraph("test paragraph")


# Save the document
doc.save('output.docx')