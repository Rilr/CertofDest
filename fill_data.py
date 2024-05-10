import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


# Read the CSV file
df = pd.read_csv('assets/codd.csv')
col_labels = df.columns

# Create a new Word document
doc = Document('assets/codd.docx')

# Iterate over each row in the dataframe
for index, row in df.iterrows():
    # Add three new lines
    doc.add_paragraph('\n\n')
    # Title Text
    titleparagraph = doc.add_paragraph()
    titlerun = titleparagraph.add_run('CERTIFICATE OF DESTRUCTION')
    titlerun.font.name = 'Georgia (Headings)'
    titlerun.font.size = Pt(26)
    titlerun.font.color.rgb = RGBColor(152, 134, 0)
    titleparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Preamble Text
    preparagraph = doc.add_paragraph()
    prerun = preparagraph.add_run('This is to certify that the following drive(s) have been securely disposed of and rendered unrecoverable using the method(s) indicated below:' '\n')
    prerun.font.name = 'Garamond'
    prerun.font.size = Pt(12)
    preparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Create a table with 1 row and 2 columns
    table = doc.add_table(rows=len(col_labels), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(col_labels):
        cell = table.cell(i, 0)
        cell.text = label
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.width = Inches(2.5)
    for i, data in enumerate(row):
        cell = table.cell(i, 1)
        cell.text = str(data)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.width = Inches(3.0)
    postparagraph = doc.add_paragraph()
    postrun = postparagraph.add_run('\n' 'All drive(s) listed have been securely disposed of after having been rendered unrecoverable using the method(s) indicated above.')
    postrun.font.name = 'Garamond'
    postrun.font.size = Pt(16)
    postparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a page break after each table
    doc.add_page_break()

# Delete the first and last pages

# Save the document
doc.save('output.docx')